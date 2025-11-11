from fastapi import FastAPI, UploadFile, Form, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import re

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"]
)

def format_abstract_section(doc: Document):
    abstract_start = None
    abstract_end = None
    keywords_text = ""

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().lower()
        if text.startswith("abstract"):
            abstract_start = i
        elif abstract_start is not None and text.startswith("keywords"):
            abstract_end = i
            keywords_text = p.text.replace("Keywords:", "").strip()
            break

    if abstract_start is None:
        return None

    abstract_paragraphs = []
    for j in range(abstract_start + 1, abstract_end or len(doc.paragraphs)):
        abstract_paragraphs.append(doc.paragraphs[j].text)
        if abstract_end and j == abstract_end - 1:
            break
    abstract_text = "\n".join(abstract_paragraphs).strip()

    for _ in range(abstract_start, (abstract_end or abstract_start + 1)):
        p = doc.paragraphs[abstract_start]
        p._element.getparent().remove(p._element)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Abstract")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    box = doc.add_paragraph()
    box.paragraph_format.left_indent = Inches(0.3)
    box.paragraph_format.right_indent = Inches(0.3)
    box.paragraph_format.line_spacing = 1.15
    box.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    abstract_run = box.add_run(abstract_text + "\n\nKeywords: ")
    abstract_run.font.name = "Times New Roman"
    abstract_run.font.size = Pt(12)
    abstract_run.font.color.rgb = RGBColor(0, 0, 0)

    shading_elm = box._element.xpath(".//w:shd")
    if not shading_elm:
        pPr = box._element.get_or_add_pPr()
        shd = pPr._new_shd()
        shd.set(qn("w:fill"), "D9D9D9")

    keyword_run = box.add_run(keywords_text or "Add your keywords here.")
    keyword_run.font.name = "Times New Roman"
    keyword_run.font.size = Pt(12)
    keyword_run.font.bold = False
    keyword_run.font.color.rgb = RGBColor(0, 0, 0)

    return keywords_text

def add_footer(doc: Document, footer_text: str):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        run = paragraph.add_run(footer_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p = paragraph._element
        pPr = p.get_or_add_pPr()
        borders = pPr.xpath(".//w:pBdr")
        if not borders:
            from docx.oxml import parse_xml
            border_xml = """
            <w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="8" w:space="1" w:color="000000"/>
            </w:pBdr>
            """
            pPr.append(parse_xml(border_xml))

def set_file_metadata(doc: Document, author_name: str, keywords: str):
    title_text = ""
    for p in doc.paragraphs:
        if p.text.strip():
            title_text = p.text.strip()
            break

    props = doc.core_properties
    props.title = title_text
    props.author = author_name
    props.subject = "Academic Research Paper"
    props.keywords = keywords or ""
    props.comments = "Processed and formatted by Globus Publication system"
    props.category = "Academic"
    props.content_status = "Final"
    props.last_modified_by = author_name

@app.post("/process")
async def process_doc(
    journal_name: str = Form(...),
    volume_details: str = Form(...),
    paper_received: str = Form(...),
    paper_accepted: str = Form(...),
    paper_published: str = Form(...),
    author_name: str = Form(...),
    corresponding_author: str = Form(...),
    email: str = Form(...),
    doi: str = Form(...),
    footer_text: str = Form(...),
    file: UploadFile = File(...)
):
    temp_input = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    temp_output = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")

    with open(temp_input.name, "wb") as f:
        f.write(await file.read())

    doc = Document(temp_input.name)

    replacements = {
        "{{JOURNAL_NAME}}": journal_name,
        "{{VOLUME_DETAILS}}": volume_details,
        "{{PAPER_RECEIVED}}": paper_received,
        "{{PAPER_ACCEPTED}}": paper_accepted,
        "{{PAPER_PUBLISHED}}": paper_published,
        "{{AUTHOR_NAME}}": author_name,
        "{{CORRESPONDING_AUTHOR}}": corresponding_author,
        "{{EMAIL}}": email,
        "{{DOI}}": doi,
        "{{FOOTER}}": footer_text,
    }

    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, val)

    keywords_detected = format_abstract_section(doc)
    add_footer(doc, footer_text)
    set_file_metadata(doc, author_name, keywords_detected)

    doc.save(temp_output.name)
    return FileResponse(temp_output.name, filename="formatted_publication.docx")
